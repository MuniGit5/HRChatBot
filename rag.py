"""Simple RAG for James Shield HR Assistant - answers only from handbook documents."""

import os
from pathlib import Path
from typing import Optional

from docx import Document
from openai import OpenAI
import chromadb
from chromadb.utils import embedding_functions


def get_document_list():
    """Return list of document names (without .docx) for filtering."""
    data_dir = Path(__file__).parent / "data"
    docs = set()
    for folder in ["hr_docs", "hr_extra"]:
        folder_path = data_dir / folder
        if not folder_path.exists():
            continue
        for docx_file in folder_path.glob("*.docx"):
            if docx_file.name.startswith("~$"):
                continue
            docs.add(docx_file.name)
    return sorted(docs)


def load_documents():
    """Load all docx files from data folder."""
    data_dir = Path(__file__).parent / "data"
    texts = []
    sources = []

    for folder in ["hr_docs", "hr_extra"]:
        folder_path = data_dir / folder
        if not folder_path.exists():
            continue
        for docx_file in folder_path.glob("*.docx"):
            if docx_file.name.startswith("~$"):
                continue
            try:
                doc = Document(docx_file)
                full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                if full_text:
                    # Chunk by paragraphs (simple approach)
                    chunks = []
                    current = []
                    for p in doc.paragraphs:
                        if p.text.strip():
                            current.append(p.text)
                            if len("\n".join(current)) > 600:
                                chunks.append("\n".join(current))
                                current = []
                    if current:
                        chunks.append("\n".join(current))
                    for chunk in chunks:
                        texts.append(chunk)
                        sources.append(docx_file.name)
            except Exception as e:
                print(f"Error loading {docx_file}: {e}")

    return texts, sources


def build_vector_store():
    """Build ChromaDB vector store from documents."""
    texts, sources = load_documents()
    if not texts:
        raise ValueError("No documents loaded")

    client = chromadb.PersistentClient(path=str(Path(__file__).parent / "chroma_db"))
    ef = embedding_functions.OpenAIEmbeddingFunction(
        api_key=os.getenv("OPENAI_API_KEY"),
        model_name="text-embedding-3-small",
    )

    try:
        client.delete_collection("handbook")
    except Exception:
        pass

    collection = client.create_collection(
        name="handbook",
        embedding_function=ef,
        metadata={"hnsw:space": "cosine"},
    )

    collection.add(
        documents=texts,
        ids=[f"doc_{i}" for i in range(len(texts))],
        metadatas=[{"source": s} for s in sources],
    )
    return collection


def get_answer(question: str, document_filter: Optional[str] = None) -> str:
    """Get RAG answer - only from handbook, summarize to 2 sentences.
    If document_filter is provided (e.g. 'Vacation Policy.docx'), search only that document."""
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        return "Error: OPENAI_API_KEY not set. Please add it to .env file."

    client = chromadb.PersistentClient(path=str(Path(__file__).parent / "chroma_db"))
    ef = embedding_functions.OpenAIEmbeddingFunction(
        api_key=api_key,
        model_name="text-embedding-3-small",
    )

    try:
        collection = client.get_collection(name="handbook", embedding_function=ef)
    except Exception:
        collection = build_vector_store()

    # For PTO/accrual questions, add related terms to improve retrieval (handbook uses "Vacation" and "monthly")
    query_text = question
    if any(term in question.lower() for term in ["pto", "accrual", "accrue", "pay period"]):
        query_text = f"{question} vacation accrual monthly hours"
    query_kwargs = {"query_texts": [query_text], "n_results": 10}
    if document_filter:
        query_kwargs["where"] = {"source": document_filter}

    results = collection.query(**query_kwargs)

    if not results or not results["documents"] or not results["documents"][0]:
        return "The information you requested is confidential or privacy restricted. Please reach out to Human Resources directly for assistance."

    context = "\n\n---\n\n".join(results["documents"][0])

    openai_client = OpenAI(api_key=api_key)
    response = openai_client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {
                "role": "system",
                "content": """You are an HR assistant. Summarize the provided document context to answer the user's question.

RULES:
1. Use ONLY information from the provided context. Do NOT add, infer, or invent anything.
2. If the context contains accrual rates, hours, amounts, or policy details that relate to the question - you MUST answer using that information. Do NOT use the fallback.
3. If the context uses different terminology (e.g., "monthly" vs "pay period", "Vacation" vs "PTO"), present the information from the context. For accrual: if the handbook says "hours per month", provide those rates and note that PTO/vacation accrues monthly.
4. ONLY if the context truly has NO relevant information, respond EXACTLY with: "The information you requested is confidential or privacy restricted. Please reach out to Human Resources directly for assistance."
5. Summarize your answer to 2-3 sentences. Be concise but accurate.""",
            },
            {
                "role": "user",
                "content": f"Context from handbook:\n\n{context}\n\nQuestion: {question}",
            },
        ],
        temperature=0.1,
    )

    answer = response.choices[0].message.content.strip()
    return answer
